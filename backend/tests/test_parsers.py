import fitz
from docx import Document as DocxDocument
from openpyxl import Workbook

from app.services.parser.docx_parser import parse_docx
from app.services.parser.pdf_parser import parse_pdf
from app.services.parser.xlsx_parser import parse_xlsx


def _create_pdf(path: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Company brochure content for testing ingestion pipeline.")
    doc.save(path)
    doc.close()


def _create_docx(path: str) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Company policy document with detailed information.")
    doc.add_paragraph("Contact us at info@example.com for more details.")
    doc.save(path)


def _create_xlsx(path: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Products"
    ws.append(["Product", "Price"])
    ws.append(["Widget A", "100"])
    ws.append(["Widget B", "200"])
    wb.save(path)


def test_pdf_parser(tmp_path):
    file_path = tmp_path / "test.pdf"
    _create_pdf(str(file_path))
    sections = parse_pdf(str(file_path))
    assert len(sections) >= 1
    assert "Company brochure" in sections[0].text


def test_docx_parser(tmp_path):
    file_path = tmp_path / "test.docx"
    _create_docx(str(file_path))
    sections = parse_docx(str(file_path))
    assert len(sections) == 1
    assert "Company policy" in sections[0].text


def test_xlsx_parser(tmp_path):
    file_path = tmp_path / "test.xlsx"
    _create_xlsx(str(file_path))
    sections = parse_xlsx(str(file_path))
    assert len(sections) >= 1
    assert "Widget A" in sections[0].text
